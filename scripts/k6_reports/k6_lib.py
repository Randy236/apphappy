"""Genera informe Word y Excel de pruebas de rendimiento k6 — Happy Jump."""

from __future__ import annotations

import json
import re
import socket
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

COLOR_PRIMARIO = RGBColor(0x18, 0xB5, 0x8A)
COLOR_TEXTO = RGBColor(0x1E, 0x29, 0x3B)
HEADER_FILL = "18B58A"


@dataclass
class ServidorInfo:
    hostname: str = ""
    sistema_operativo: str = ""
    procesador: str = ""
    nucleos: str = ""
    ram_gb: str = ""
    node_version: str = ""
    k6_version: str = ""
    base_url: str = "http://localhost:3000"
    puerto_api: str = "3000"
    motor_bd: str = "MySQL 8"
    framework_api: str = "Node.js + Express"
    arquitectura: str = "Cliente Android + API REST monolítica + MySQL"


@dataclass
class MetricasK6:
    http_reqs: str = "—"
    http_req_failed: str = "—"
    http_req_duration_avg: str = "—"
    http_req_duration_med: str = "—"
    http_req_duration_p90: str = "—"
    http_req_duration_p95: str = "—"
    http_req_duration_max: str = "—"
    checks_total: str = "—"
    checks_succeeded: str = "—"
    checks_failed: str = "—"
    iterations: str = "—"
    data_received: str = "—"
    data_sent: str = "—"
    thresholds_ok: bool = True
    escenarios: list[str] = field(default_factory=list)
    checks_list: list[str] = field(default_factory=list)
    raw_log: str = ""


def _estilo_header(cell) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), HEADER_FILL)
    cell._tc.get_or_add_tcPr().append(shading)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _titulo(doc: Document, texto: str, nivel: int = 1) -> None:
    p = doc.add_heading(texto, level=nivel)
    for r in p.runs:
        r.font.color.rgb = COLOR_PRIMARIO if nivel == 1 else COLOR_TEXTO


def _tabla(doc: Document, headers: list[str], filas: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(filas), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        _estilo_header(t.rows[0].cells[i])
    for ri, fila in enumerate(filas):
        for ci, val in enumerate(fila):
            t.rows[ri + 1].cells[ci].text = str(val)


def parsear_log_k6(texto: str) -> MetricasK6:
    m = MetricasK6(raw_log=texto)
    patrones = {
        "http_reqs": r"http_reqs\.+?:\s*(.+)",
        "http_req_failed": r"http_req_failed\.+?:\s*(.+)",
        "http_req_duration_avg": r"http_req_duration\.+?:\s*avg=([^\s]+)",
        "http_req_duration_med": r"med=([^\s]+)",
        "http_req_duration_p90": r"p\(90\)=([^\s]+)",
        "http_req_duration_p95": r"p\(95\)=([^\s]+)",
        "http_req_duration_max": r"max=([^\s]+)",
        "checks_total": r"checks_total\.+?:\s*(\d+)",
        "checks_succeeded": r"checks_succeeded\.+?:\s*([\d.]+%)",
        "checks_failed": r"checks_failed\.+?:\s*([\d.]+%)",
        "iterations": r"iterations\.+?:\s*(\d+)",
        "data_received": r"data_received\.+?:\s*(.+)",
        "data_sent": r"data_sent\.+?:\s*(.+)",
    }
    for key, pat in patrones.items():
        match = re.search(pat, texto)
        if match:
            setattr(m, key, match.group(1).strip())
    m.checks_list = re.findall(r"✓\s+(.+)", texto.replace("Ô£ô", "✓"))
    if not m.checks_list:
        m.checks_list = re.findall(r"checks_succeeded.*\n((?:.*\n)*)", texto)[:5]  # fallback
    m.thresholds_ok = "thresholds on metrics 'http_req_failed' have been crossed" not in texto.lower()
    if "✗" in texto or "thresholds" in texto.lower() and "crossed" in texto.lower():
        m.thresholds_ok = False
    for esc in ("smoke", "load", "stress"):
        if esc in texto.lower():
            m.escenarios.append(esc)
    if not m.escenarios:
        m.escenarios = ["smoke", "load", "stress"]
    return m


def parsear_summary_json(ruta: Path) -> dict[str, Any]:
    if not ruta.exists():
        return {}
    return json.loads(ruta.read_text(encoding="utf-8"))


def generar_word(servidor: ServidorInfo, metricas: MetricasK6, salida: Path) -> Path:
    doc = Document()
    fecha = datetime.now()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("INFORME DE PRUEBAS DE RENDIMIENTO")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = COLOR_PRIMARIO

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Happy Jump — API REST\n").bold = True
    sub.add_run(f"Fecha: {fecha.strftime('%d/%m/%Y %H:%M')}\n")
    sub.add_run(f"URL objetivo: {servidor.base_url}\n")
    sub.add_run("Herramienta: Grafana k6 (pruebas de carga y estrés)")

    doc.add_paragraph()

    _titulo(doc, "1. Introducción", 2)
    doc.add_paragraph(
        "Este informe documenta las pruebas de rendimiento ejecutadas sobre la API "
        "Happy Jump, considerando las características del servidor de pruebas y los "
        "escenarios smoke, carga (load) y estrés (stress) definidos en k6."
    )

    _titulo(doc, "2. Características del servidor de pruebas", 2)
    doc.add_paragraph(
        "El rendimiento medido depende del hardware y software donde corre la API. "
        "A continuación se detalla el entorno utilizado en esta ejecución:"
    )
    _tabla(
        doc,
        ["Característica", "Valor"],
        [
            ["Nombre del equipo", servidor.hostname or "—"],
            ["Sistema operativo", servidor.sistema_operativo or "—"],
            ["Procesador", servidor.procesador or "—"],
            ["Núcleos lógicos", servidor.nucleos or "—"],
            ["Memoria RAM", f"{servidor.ram_gb} GB" if servidor.ram_gb else "—"],
            ["Node.js", servidor.node_version or "—"],
            ["k6", servidor.k6_version or "—"],
            ["Arquitectura del sistema", servidor.arquitectura],
            ["Framework API", servidor.framework_api],
            ["Base de datos", servidor.motor_bd],
            ["Puerto API", servidor.puerto_api],
            ["URL base", servidor.base_url],
        ],
    )
    doc.add_paragraph(
        "Nota: la API es un proceso monolítico Express (un solo servicio Node.js) "
        "que escucha en 0.0.0.0:3000 y se conecta a MySQL en localhost. No hay "
        "balanceador de carga ni contenedor Docker en el entorno de desarrollo local."
    )

    _titulo(doc, "3. Objetivos de las pruebas", 2)
    for obj in [
        "Medir latencia y tasa de error bajo carga concurrente simulada.",
        "Validar que los endpoints críticos (health, login, reservas, reportes) responden dentro de umbrales aceptables.",
        "Identificar el comportamiento del servidor al incrementar usuarios virtuales (VUs).",
        "Documentar métricas para comparación en futuras versiones.",
    ]:
        doc.add_paragraph(obj, style="List Bullet")

    _titulo(doc, "4. Metodología y escenarios k6", 2)
    doc.add_paragraph("Script: k6/performance-api.js")
    _tabla(
        doc,
        ["Escenario", "Executor", "Descripción"],
        [
            [
                "Smoke",
                "1 VU, 15 s",
                "Verificación básica: /health, login, listar cancha y reportes.",
            ],
            [
                "Carga (load)",
                "0→10→20 VUs (2 min)",
                "Flujo típico: health, login, consultas de reservas y reportes.",
            ],
            [
                "Estrés (stress)",
                "0→30→60 VUs (1 min)",
                "Pico sobre endpoints ligeros: /health y /sumar.",
            ],
        ],
    )
    doc.add_paragraph("Umbrales configurados (thresholds):")
    for u in [
        "Tasa de error HTTP < 5 %",
        "Percentil 95 de duración < 800 ms",
        "Checks exitosos ≥ 90 %",
    ]:
        doc.add_paragraph(u, style="List Bullet")

    _titulo(doc, "5. Resultados de la ejecución", 2)
    estado = "CUMPLIDO" if metricas.thresholds_ok else "REVISAR"
    doc.add_paragraph(f"Estado global de umbrales: {estado}")
    _tabla(
        doc,
        ["Métrica", "Valor"],
        [
            ["Peticiones HTTP totales", metricas.http_reqs],
            ["Tasa de fallo HTTP", metricas.http_req_failed],
            ["Duración promedio", metricas.http_req_duration_avg],
            ["Duración mediana", metricas.http_req_duration_med],
            ["Percentil 90", metricas.http_req_duration_p90],
            ["Percentil 95", metricas.http_req_duration_p95],
            ["Duración máxima", metricas.http_req_duration_max],
            ["Checks totales", metricas.checks_total],
            ["Checks exitosos", metricas.checks_succeeded],
            ["Checks fallidos", metricas.checks_failed],
            ["Iteraciones", metricas.iterations],
            ["Datos recibidos", metricas.data_received],
            ["Datos enviados", metricas.data_sent],
        ],
    )

    if metricas.checks_list:
        _titulo(doc, "5.1 Checks individuales", 3)
        for c in metricas.checks_list[:25]:
            doc.add_paragraph(f"✓ {c}", style="List Bullet")

    _titulo(doc, "6. Análisis considerando el servidor", 2)
    doc.add_paragraph(
        "En un servidor de desarrollo local (Windows + Node.js + MySQL en la misma máquina), "
        "los tiempos de respuesta incluyen latencia de red loopback, procesamiento JavaScript "
        "en un solo hilo de event loop de Node.js, y consultas SQL a MySQL."
    )
    doc.add_paragraph(
        "Bajo carga moderada (20 VUs), el cuello de botella típico es la concurrencia de "
        "consultas a la base de datos y el login con validación JWT + bcrypt. Los endpoints "
        "de solo lectura (/health, GET reservas) suelen responder más rápido que POST con escritura."
    )
    doc.add_paragraph(
        "En el escenario de estrés (hasta 60 VUs sobre /health), se observa el límite del "
        "entorno local. En producción se recomienda: más RAM, pool de conexiones MySQL ajustado, "
        "HTTPS, y despliegue en servidor dedicado o contenedor con límites de recursos definidos."
    )

    _titulo(doc, "7. Criterios de aceptación", 2)
    _tabla(
        doc,
        ["Criterio", "Umbral", "Resultado"],
        [
            ["Error HTTP", "< 5 %", "Ver métrica http_req_failed"],
            ["Latencia p95", "< 800 ms", metricas.http_req_duration_p95],
            ["Checks", "≥ 90 % OK", metricas.checks_succeeded],
            ["API disponible bajo carga", "Sin caídas", estado],
        ],
    )

    _titulo(doc, "8. Conclusiones", 2)
    if metricas.thresholds_ok:
        doc.add_paragraph(
            "Las pruebas de rendimiento con k6 sobre la API Happy Jump se ejecutaron "
            "correctamente. Los umbrales definidos se cumplieron en el entorno de pruebas "
            "documentado. El sistema soporta la carga simulada para el uso esperado en un "
            "centro de entretenimiento con varios dispositivos móviles concurrentes."
        )
    else:
        doc.add_paragraph(
            "Se detectaron desviaciones en los umbrales de rendimiento. Se recomienda revisar "
            "sesiones bloqueadas (active_token), conexión MySQL, y repetir tras liberar sesiones "
            "con: node server/scripts/clear-sessions-for-k6.mjs"
        )
    doc.add_paragraph(
        "Recomendaciones: ejecutar este informe antes de cada entrega; comparar métricas "
        "entre versiones; en producción usar servidor con recursos superiores al PC de desarrollo."
    )

    _titulo(doc, "9. Anexos", 2)
    doc.add_paragraph("A. k6/performance-api.js — Script de pruebas")
    doc.add_paragraph("B. k6/results/ — Salida de consola y summary JSON")
    doc.add_paragraph("C. docs/GUIA_ITEM6_K6_TODOS_ENDPOINTS.md — Guía k6 del proyecto")

    salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(salida)
    return salida


def generar_excel(servidor: ServidorInfo, metricas: MetricasK6, salida: Path) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Resumen"
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor=HEADER_FILL)

    filas = [
        ("Proyecto", "Happy Jump"),
        ("Fecha", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("URL API", servidor.base_url),
        ("SO", servidor.sistema_operativo),
        ("Node.js", servidor.node_version),
        ("k6", servidor.k6_version),
        ("RAM (GB)", servidor.ram_gb),
        ("Núcleos", servidor.nucleos),
        ("http_reqs", metricas.http_reqs),
        ("http_req_failed", metricas.http_req_failed),
        ("avg", metricas.http_req_duration_avg),
        ("p95", metricas.http_req_duration_p95),
        ("checks OK", metricas.checks_succeeded),
        ("Umbrales", "OK" if metricas.thresholds_ok else "REVISAR"),
    ]
    ws.append(["Campo", "Valor"])
    for c in ws[1]:
        c.font = header_font
        c.fill = header_fill
    for campo, valor in filas:
        ws.append([campo, valor])
    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 50

    ws2 = wb.create_sheet("Escenarios")
    ws2.append(["Escenario", "VUs", "Duración", "Endpoints"])
    for c in ws2[1]:
        c.font = header_font
        c.fill = header_fill
    esc_data = [
        ("Smoke", "1", "15 s", "/health, login, reservas, reportes"),
        ("Carga", "10→20", "2 min", "Flujo lectura autenticado"),
        ("Estrés", "30→60", "1 min", "/health, /sumar"),
    ]
    for row in esc_data:
        ws2.append(list(row))

    salida.parent.mkdir(parents=True, exist_ok=True)
    wb.save(salida)
    return salida


def cargar_servidor_desde_json(ruta: Path) -> ServidorInfo:
    if not ruta.exists():
        return ServidorInfo(hostname=socket.gethostname())
    data = json.loads(ruta.read_text(encoding="utf-8-sig"))
    return ServidorInfo(**{k: v for k, v in data.items() if k in ServidorInfo.__dataclass_fields__})
