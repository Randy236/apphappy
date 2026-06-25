"""Checklist SDLC Audit — Happy Jump. Genera Word y Excel."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

HEADER_FILL = "18B58A"
OK_FILL = "DCFCE7"
NO_FILL = "FEE2E2"


@dataclass
class ItemChecklist:
    codigo: str
    area: str
    item: str
    cumple: bool
    observaciones: str


def _items() -> list[ItemChecklist]:
    return [
        # METODOLOGÍA SDLC
        ItemChecklist("CL-01-01", "Metodología SDLC",
            "¿Se han determinado el alcance de las responsabilidades de la dirección, auditoría interna, usuarios, QA y procesamiento de datos durante el diseño, desarrollo y mantenimiento del sistema?",
            True, "Roles definidos en docs/SCRUM_CMMI_HAPPY_JUMP.md (Product Owner, Scrum Master, Dev Team). QA: pruebas Postman, Sonar, ZAP, k6. Usuario final: trabajador/admin vía app Android."),
        ItemChecklist("CL-01-02", "Metodología SDLC",
            "¿Los workpapers del SDLC evidencian que se obtuvieron los niveles de autorización apropiados para cada fase?",
            True, "Evidencia: entregables por unidad (entregableunidad/entregable-01…12), commits en GitHub, GitHub Actions y Jenkinsfile. Actas formales por fase académica en docs/DOCUMENTO_AVANCE_UNIDAD2.md."),
        ItemChecklist("CL-01-03", "Metodología SDLC",
            "¿Existe una metodología formal de desarrollo implantada y soportada por herramientas CASE?",
            False, "Metodología formal: Scrum + alineación CMMI Nivel 2 (docs/SCRUM_CMMI_HAPPY_JUMP.md). No se usan herramientas CASE comerciales; se emplean Jira, GitHub, Android Studio, Postman, SonarCloud como soporte."),
        ItemChecklist("CL-01-04", "Metodología SDLC",
            "¿El proyecto tiene un plan formal de proyecto documentado?",
            True, "Plan documentado: docs/SCRUM_CMMI_HAPPY_JUMP.md, docs/GUIA_ITEM1_DOCUMENTO_AVANCE.md, sprints 1–3, cronograma en docs y backlog Jira (docs/jira-import/happy-jump-backlog.csv)."),
        ItemChecklist("CL-01-05", "Metodología SDLC",
            "¿El proyecto tiene definición de alcance y marco formal?",
            True, "Alcance: README.md (módulos Cancha, Salones, Reportes, Perfil, Auth). Marco: ISO/IEC 12207, ISO/IEC 25010, CMMI-DEV Nivel 2 referenciados en Project Charter y SCRUM_CMMI."),
        ItemChecklist("CL-01-06", "Metodología SDLC",
            "¿Están especificados los entregables principales, plazos y roles/responsabilidades?",
            True, "Entregables: entregableunidad/ (01–12). Roles: tabla en SCRUM_CMMI. Plazos: sprints de 2 semanas y fechas en Project Charter de auditoría."),
        ItemChecklist("CL-01-07", "Metodología SDLC",
            "¿Se realiza análisis de riesgos del proyecto?",
            False, "No existe matriz de riesgos formal del proyecto de desarrollo. Riesgos identificados de forma informal (MySQL apagado, sesiones bloqueadas, red Wi‑Fi). Recomendación: crear docs/MATRIZ_RIESGOS.md."),
        ItemChecklist("CL-01-08", "Metodología SDLC",
            "¿Se están siguiendo los procedimientos definidos para el área de desarrollo?",
            True, "Procedimientos seguidos: Definition of Done (SCRUM_CMMI), npm run test:ci, Gradle testDebugUnitTest, flujo Git (feat/fix/docs), CI en .github/workflows/."),
        # ANÁLISIS DE NECESIDADES
        ItemChecklist("CL-02-01", "Análisis de necesidades",
            "¿Existen procedimientos formales para realizar el análisis de necesidades?",
            True, "Catálogo docs/REQUERIMIENTOS.csv (REQ-001 a REQ-013) con prioridad, módulo y criterio de aceptación. Backlog Jira importable en docs/jira-import/."),
        ItemChecklist("CL-02-02", "Análisis de necesidades",
            "¿El análisis de necesidades de un proyecto reciente cumple con los estándares establecidos?",
            True, "Requisitos alineados a módulos implementados: login, cancha, salones, reportes, perfil, validaciones API. Criterios verificables en Postman y casos CP-XXX."),
        ItemChecklist("CL-02-03", "Análisis de necesidades",
            "¿Existe un mecanismo para registrar necesidades de desarrollo con descripción, riesgos y análisis coste/beneficio?",
            False, "Jira y REQUERIMIENTOS.csv registran descripción y prioridad. No hay análisis coste/beneficio formal documentado."),
        ItemChecklist("CL-02-04", "Análisis de necesidades",
            "¿Qué tipo de especificaciones de requisitos se están empleando?",
            True, "User stories en Jira, requisitos funcionales en CSV (REQ-XXX), casos de prueba CP-XXX en docs/CASOS_DE_PRUEBA.csv, contratos API en server/openapi.json."),
        ItemChecklist("CL-02-05", "Análisis de necesidades",
            "¿Dónde se almacenan los requisitos? ¿Se usa una herramienta estándar?",
            True, "Almacenamiento: docs/REQUERIMIENTOS.csv, docs/CASOS_DE_PRUEBA.csv, Jira (backlog CSV), OpenAPI en server/openapi.json."),
        ItemChecklist("CL-02-06", "Análisis de necesidades",
            "¿Existe un proceso de revisión de requisitos?",
            True, "Revisión vía Sprint Planning, refinamiento de backlog Jira, validación con pruebas Postman (VAL-XXX) y casos CP en entregables 6–7."),
        ItemChecklist("CL-02-07", "Análisis de necesidades",
            "¿Se gestiona la trazabilidad de requisitos?",
            True, "Trazabilidad parcial documentada: REQ → módulo app/API → CP-XXX en CASOS_DE_PRUEBA.csv → tests en app/src/test/ y server/test/."),
        # DISEÑO Y DESARROLLO
        ItemChecklist("CL-03-01", "Diseño y desarrollo",
            "¿Se han revisado las especificaciones de diseño y hay evidencia escrita de aprobación?",
            True, "Diseño: server/openapi.json (Swagger UI), docs/DOCUMENTO_AVANCE_UNIDAD2.md, arquitectura cliente-servidor en README. Aprobación académica por entregables de unidad."),
        ItemChecklist("CL-03-02", "Diseño y desarrollo",
            "¿Las especificaciones de diseño cumplen con los estándares?",
            True, "OpenAPI 3 documenta endpoints REST. App Android con capas data/ui/viewmodel. Dominio separado en server/src/domain/."),
        ItemChecklist("CL-03-03", "Diseño y desarrollo",
            "¿Se incorporan pista de auditoría y controles programados en las especificaciones de diseño?",
            True, "Migración 010_auditoria.sql, server/src/domain/auditoria.js, endpoint GET /auditoria (admin). Eliminado lógico: migration 009, softDelete.js."),
        ItemChecklist("CL-03-04", "Diseño y desarrollo",
            "¿Los documentos fuente para entrada de datos están diseñados para facilitar la captura precisa?",
            True, "Formularios Android (LoginScreen, CanchaScreen, SalonesScreen) con validación UI. API: server/src/domain/entradaValidacion.js + tests."),
        ItemChecklist("CL-03-05", "Diseño y desarrollo",
            "¿Los programas cumplen con los estándares de programación del área?",
            True, "Convenciones Git (feat/fix/test/docs). Análisis SonarCloud. Kotlin en app/, ES modules en server/. Exclusiones UI documentadas en sonar-project.properties."),
        ItemChecklist("CL-03-06", "Diseño y desarrollo",
            "¿Existen estándares documentados de codificación en una wiki colaborativa?",
            False, "Estándares en README.md y docs/ del repositorio GitHub, no en wiki externa (Confluence). Equivalente funcional cumplido vía repo."),
        ItemChecklist("CL-03-07", "Diseño y desarrollo",
            "¿El equipo trabaja con un diseñador desde el inicio del proyecto?",
            False, "UI diseñada directamente en Jetpack Compose por el equipo de desarrollo. No hay diseñador UX/UI externo documentado."),
        ItemChecklist("CL-03-08", "Diseño y desarrollo",
            "¿Se utiliza un sistema de control de versiones conforme a las mejores prácticas?",
            True, "Git + GitHub (https://github.com/Randy236/apphappy). Ramas main/develop, .gitignore, historial de commits, PR template .github/pull_request_template.md."),
        ItemChecklist("CL-03-09", "Diseño y desarrollo",
            "¿Existe integración continua / despliegue continuo (CI/CD)?",
            True, "GitHub Actions: api-tests.yml, sonarcloud.yml, snyk.yml. Jenkinsfile (build, Sonar, deploy). Scripts ci/ y docs/GUIA_JENKINS_CI_CD.md."),
        # PROCEDIMIENTOS DE PRUEBA
        ItemChecklist("CL-04-01", "Procedimientos de prueba",
            "¿Existen procedimientos documentados de prueba de sistemas y programas?",
            True, "docs/CASOS_DE_PRUEBA.csv, docs/informes-item6/INFORME_PRUEBAS_SISTEMA_HAPPY_JUMP.md, GUIA_ITEM5/6/8, scripts Postman y k6."),
        ItemChecklist("CL-04-02", "Procedimientos de prueba",
            "¿Los procedimientos de prueba, datos de prueba y resultados son comprensivos y siguen los estándares?",
            True, "Usuarios seed Admin/Rosisela PIN 1234. Resultados: entregable-02 (Postman), entregable-08 (ZAP), entregable8.1 (k6), RESULTADOS_EJECUCION.csv."),
        ItemChecklist("CL-04-03", "Procedimientos de prueba",
            "¿Son adecuadas las pruebas realizadas sobre las fases manuales de la aplicación?",
            True, "Pruebas E2E manuales CP-001–CP-022 documentadas. Manual de usuario entregable-09. Capturas en informes de sistema."),
        ItemChecklist("CL-04-04", "Procedimientos de prueba",
            "¿Existe una estrategia de pruebas documentada?",
            True, "Estrategia en docs/informes-item6 y CHECKLIST: unitarias (JVM+Node), integración (Postman), E2E (manual CP), seguridad (ZAP+SEC), rendimiento (k6), calidad estática (Sonar/Snyk)."),
        ItemChecklist("CL-04-05", "Procedimientos de prueba",
            "¿Hay proceso de revisión de calidad y métricas de software?",
            True, "SonarCloud cobertura ≥80% (entregable-03). JaCoCo XML. Snyk dependencias. Métricas en METRICAS_SONAR_HAPPY_JUMP.xlsx."),
        ItemChecklist("CL-04-06", "Procedimientos de prueba",
            "¿Se gestiona la integración de software y la documentación de pruebas?",
            True, "Postman collections en docs/postman/. Scripts generan Word/Excel (scripts/postman_reports/). Allure: scripts/run-allure-report.ps1."),
        ItemChecklist("CL-04-07", "Procedimientos de prueba",
            "¿Existe un sistema de gestión de defectos?",
            True, "Gestión de defectos vía issues Jira (backlog CSV) y GitHub Issues. Severidad en informes ZAP y Matriz de Hallazgos (auditoría)."),
        ItemChecklist("CL-04-08", "Procedimientos de prueba",
            "¿Cómo funciona el proceso de revisión de código?",
            True, "Pull Requests en GitHub con template. Guía docs/GUIA_ITEM3_GITHUB_REVIEWS_COLABORATIVO.md. SonarCloud en CI como gate de calidad."),
        # IMPLEMENTACIÓN
        ItemChecklist("CL-05-01", "Implementación",
            "¿Existen procedimientos formales de promoción e implementación de programas?",
            True, "ci/build-production.sh, ci/deploy-server.sh, docs/GUIA_JENKINS_CI_CD.md, infra/docker/docker-compose.yml."),
        ItemChecklist("CL-05-02", "Implementación",
            "¿La documentación del procedimiento de promoción muestra que los estándares se siguen?",
            True, "Jenkinsfile define stages: checkout → test → Sonar → build → deploy. Evidencias en docs/evidencias-jenkins/."),
        ItemChecklist("CL-05-03", "Implementación",
            "¿Los cambios seleccionados tienen registros de soporte que evidencian aprobación adecuada?",
            True, "Historial Git con mensajes convencionales. PR reviews. Migraciones numeradas server/migrations/001–010 con scripts apply-*.mjs."),
        ItemChecklist("CL-05-04", "Implementación",
            "¿La documentación de la implementación de nuevas aplicaciones muestra que se siguieron los procedimientos?",
            True, "README sección CI/CD. COMO_EJECUTAR.txt por entregable. local.properties.example para configuración Android."),
        ItemChecklist("CL-05-05", "Implementación",
            "¿Existe un proceso de gestión de cambios con autorización formal?",
            True, "Flujo Git: develop → main. PR obligatorio documentado. Control de cambios en Jira (historias Done)."),
        ItemChecklist("CL-05-06", "Implementación",
            "¿Cómo funciona el workflow configurado para el despliegue?",
            True, "Jenkins pipeline + docker-compose. API npm start puerto 3000. APK via Gradle installDebug. Deploy script a ~/servers/happyjump/."),
        # POST-IMPLEMENTACIÓN
        ItemChecklist("CL-06-01", "Revisión post-implementación",
            "¿Existen procedimientos formales de revisión post-implementación?",
            False, "Sprint Retrospective documentada en plantilla docs/scrum/plantilla-retrospectiva.md. No hay procedimiento formal post-producción (sistema académico/local)."),
        ItemChecklist("CL-06-02", "Revisión post-implementación",
            "¿Las modificaciones de programas, procedimientos de prueba y documentación de soporte siguen los estándares?",
            True, "Cambios incluyen tests (npm run test:ci, Gradle tests). Migraciones SQL versionadas. Docs actualizados en mismo commit/entregable."),
        ItemChecklist("CL-06-03", "Revisión post-implementación",
            "¿Se documentan y gestionan las lecciones aprendidas?",
            True, "Retrospectivas Scrum: docs/scrum/plantilla-retrospectiva.md, evidencias-scrum-cmmi/. Lecciones en DOCUMENTO_AVANCE_UNIDAD2.md."),
        ItemChecklist("CL-06-04", "Revisión post-implementación",
            "¿Se realiza seguimiento de los objetivos del sistema tras la implementación?",
            False, "Seguimiento académico por entregables e informes Sonar/k6. No hay monitoreo APM en producción (entorno local/Laragon)."),
        # MANTENIMIENTO
        ItemChecklist("CL-07-01", "Mantenimiento",
            "¿Existen procedimientos formales de mantenimiento de aplicaciones?",
            True, "Migraciones SQL incrementales. Scripts mantenimiento: clear-sessions-for-k6.mjs, demo:reset-ventas. README actualización dependencias."),
        ItemChecklist("CL-07-02", "Mantenimiento",
            "¿Las modificaciones de programas, pruebas y documentación siguen los estándares?",
            True, "Cada cambio validado con tests automatizados. Soft delete y auditoría como mantenimiento evolutivo documentado."),
        ItemChecklist("CL-07-03", "Mantenimiento",
            "¿Se gestiona el mantenimiento evolutivo y correctivo de manera diferenciada?",
            False, "Commits feat/fix distinguen evolutivo/correctivo en Git. No hay procedimiento formal escrito que diferencie ambos tipos."),
        ItemChecklist("CL-07-04", "Mantenimiento",
            "¿Existe un catálogo de componentes software reutilizables accesible y actualizado?",
            False, "Módulos reutilizables en server/src/domain/ y ui/util/ pero sin catálogo formal de componentes publicado."),
        # SOFTWARE DE SISTEMA
        ItemChecklist("CL-08-01", "Software de sistema",
            "¿Existen procedimientos formales de modificación de software de sistema?",
            False, "Stack documentado (Node, MySQL, JDK, k6) en README. No hay procedimiento formal de cambio de versiones de plataforma."),
        ItemChecklist("CL-08-02", "Software de sistema",
            "¿Las modificaciones de software de sistema tienen pruebas y documentación de soporte siguiendo los estándares?",
            True, "Actualizaciones de dependencias verificadas con npm test y Gradle test. Snyk escanea vulnerabilidades en CI."),
        ItemChecklist("CL-08-03", "Software de sistema",
            "¿Existe documentación del software de sistema desarrollado internamente y de las características del software propietario?",
            True, "README: JDK 11+, Node 18+, MySQL 8+, Android SDK. server/package.json y gradle/libs.versions.toml versionan dependencias."),
        ItemChecklist("CL-08-04", "Software de sistema",
            "¿Los lenguajes, compiladores y herramientas CASE han sido previamente homologados?",
            False, "Herramientas estándar de industria (Kotlin, Node, k6, ZAP) sin acta formal de homologación institucional."),
        # DOCUMENTACIÓN
        ItemChecklist("CL-09-01", "Estándares de documentación",
            "¿Los estándares de documentación son completos y cubren todos los artefactos del SDLC?",
            True, "docs/ cubre requisitos, pruebas, CI/CD, Scrum/CMMI, Swagger, seguridad, rendimiento, manual usuario, entregables 01–12."),
        ItemChecklist("CL-09-02", "Estándares de documentación",
            "¿Existe un estándar general para documentación técnica (análisis, diseño, programas, cuadernos de carga)?",
            True, "OpenAPI (diseño API), REQUERIMIENTOS.csv (análisis), README y DOCUMENTO_AVANCE (técnico), domain/ (lógica documentada por tests)."),
        ItemChecklist("CL-09-03", "Estándares de documentación",
            "¿Existe un estándar para manuales de usuario y procedimientos de operación?",
            True, "docs/informes-item9/MANUAL_USUARIO_HAPPY_JUMP.md, entregable-09/manualdeusuario.pdf, COMO_EJECUTAR.txt por entregable."),
        ItemChecklist("CL-09-04", "Estándares de documentación",
            "¿Los estándares son conocidos y respetados en el área?",
            True, "Equipo sigue convenciones de commits, DoD en SCRUM_CMMI, guías GUIA_ITEM*.md usadas en cada entregable académico."),
        ItemChecklist("CL-09-05", "Estándares de documentación",
            "¿Las modificaciones a estándares se difunden oportunamente dentro del área?",
            True, "Actualizaciones vía Git push al repositorio compartido. README y guías actualizados en commits docs:."),
        ItemChecklist("CL-09-06", "Estándares de documentación",
            "¿La documentación permite la trazabilidad completa a lo largo del ciclo de vida?",
            True, "Trazabilidad REQ → Jira → código → CP-XXX → informes → entregables. OpenAPI alineado a server/src/index.js (verify-swagger-coverage.mjs)."),
    ]


def _header_cell(cell) -> None:
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), HEADER_FILL)
    cell._tc.get_or_add_tcPr().append(sh)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def generar_word(items: list[ItemChecklist], salida: Path) -> Path:
    doc = Document()
    fecha = datetime.now()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LISTA DE VERIFICACIÓN SDLC\n(SDLC AUDIT CHECKLIST)")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x18, 0xB5, 0x8A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Proyecto: Happy Jump\n").bold = True
    sub.add_run(f"Código: AUD-SDLC-HAPPY-JUMP-2026-001\nFecha: {fecha.strftime('%d/%m/%Y')}\n")
    sub.add_run("Basado en System Development Life Cycle Audit Checklist estándar")

    doc.add_paragraph(
        "Instrumento operativo de recolección de evidencia para cada objetivo de auditoría. "
        "Marcar ☑ en Cumple o No Cumple según columna Estado."
    )

    area_actual = ""
    for item in items:
        if item.area != area_actual:
            area_actual = item.area
            doc.add_heading(area_actual.upper(), level=2)

        tbl = doc.add_table(rows=2, cols=4)
        tbl.style = "Table Grid"
        hdr = ["Código", "Estado", "Cumple", "No Cumple"]
        for i, h in enumerate(hdr):
            tbl.rows[0].cells[i].text = h
            _header_cell(tbl.rows[0].cells[i])

        estado = "CUMPLE" if item.cumple else "NO CUMPLE"
        tbl.rows[1].cells[0].text = item.codigo
        tbl.rows[1].cells[1].text = estado
        tbl.rows[1].cells[2].text = "☑" if item.cumple else "☐"
        tbl.rows[1].cells[3].text = "☐" if item.cumple else "☑"

        p = doc.add_paragraph()
        p.add_run("Item: ").bold = True
        p.add_run(item.item)

        p2 = doc.add_paragraph()
        p2.add_run("Observaciones: ").bold = True
        p2.add_run(item.observaciones)

        doc.add_paragraph()

    # Resumen
    total = len(items)
    ok = sum(1 for i in items if i.cumple)
    doc.add_heading("RESUMEN DE CUMPLIMIENTO", level=1)
    doc.add_paragraph(f"Total ítems evaluados: {total}")
    doc.add_paragraph(f"Cumple: {ok} ({100*ok/total:.1f}%)")
    doc.add_paragraph(f"No cumple: {total-ok} ({100*(total-ok)/total:.1f}%)")
    doc.add_paragraph(
        "Los ítems marcados como No Cumple tienen observaciones con acción recomendada o "
        "equivalente parcial documentado en el repositorio Happy Jump."
    )

    salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(salida)
    return salida


def generar_excel(items: list[ItemChecklist], salida: Path) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Checklist SDLC"

    headers = ["Código", "Área", "Item de verificación", "Cumple", "No Cumple", "Observaciones"]
    ws.append(headers)
    hf = Font(bold=True, color="FFFFFF")
    hfill = PatternFill("solid", fgColor=HEADER_FILL)
    okf = PatternFill("solid", fgColor=OK_FILL)
    nof = PatternFill("solid", fgColor=NO_FILL)

    for c in ws[1]:
        c.font = hf
        c.fill = hfill

    for item in items:
        ws.append([
            item.codigo,
            item.area,
            item.item,
            "X" if item.cumple else "",
            "" if item.cumple else "X",
            item.observaciones,
        ])
        row = ws.max_row
        fill = okf if item.cumple else nof
        ws.cell(row, 4).fill = fill
        ws.cell(row, 5).fill = fill

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["C"].width = 55
    ws.column_dimensions["D"].width = 10
    ws.column_dimensions["E"].width = 12
    ws.column_dimensions["F"].width = 60
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    # Resumen
    ws2 = wb.create_sheet("Resumen")
    ok = sum(1 for i in items if i.cumple)
    ws2.append(["Métrica", "Valor"])
    ws2.append(["Proyecto", "Happy Jump"])
    ws2.append(["Fecha", datetime.now().strftime("%Y-%m-%d")])
    ws2.append(["Total ítems", len(items)])
    ws2.append(["Cumple", ok])
    ws2.append(["No cumple", len(items) - ok])
    ws2.append(["% Cumplimiento", f"{100*ok/len(items):.1f}%"])

    salida.parent.mkdir(parents=True, exist_ok=True)
    wb.save(salida)
    return salida
