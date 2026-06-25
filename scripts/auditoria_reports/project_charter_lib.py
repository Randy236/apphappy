"""Genera Project Charter de Auditoria SDLC — Happy Jump (Word)."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

COLOR_PRIMARIO = RGBColor(0x18, 0xB5, 0x8A)
COLOR_TEXTO = RGBColor(0x1E, 0x29, 0x3B)
HEADER_FILL = "18B58A"


def _header_cell(cell) -> None:
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), HEADER_FILL)
    cell._tc.get_or_add_tcPr().append(sh)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _titulo(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = COLOR_PRIMARIO if level == 1 else COLOR_TEXTO


def _tabla(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        _header_cell(t.rows[0].cells[i])
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def generar_charter(salida: Path) -> Path:
    doc = Document()
    fecha = datetime.now()

    # Portada
    for text, size, bold in [
        ("UNIVERSIDAD PERUANA UNIÓN", 14, True),
        ("INGENIERÍA DE SISTEMAS", 12, True),
        ("Ciclo VII", 11, False),
        ("", 6, False),
        ("PROJECT CHARTER DE AUDITORÍA DEL CICLO DE VIDA", 14, True),
        ("DEL DESARROLLO DE SOFTWARE (SDLC)", 14, True),
        ("", 6, False),
        ("Profesor responsable: Ing. Rubén Roque Sucari", 11, False),
        ("", 8, False),
        ("PROYECTO HAPPY JUMP", 16, True),
        ("Sistema de reservas de cancha y salones", 12, False),
        ("", 6, False),
        ("Nombre del producto: Happy Jump", 11, False),
        ("Repositorio: https://github.com/Randy236/apphappy", 10, False),
        ("Grupo: [Completar nombre del grupo]", 11, False),
        ("Integrantes: [Nombres y apellidos — Código]", 11, False),
        (f"Juliaca, {fecha.strftime('%B %Y').replace('January','enero').replace('February','febrero').replace('March','marzo').replace('April','abril').replace('May','mayo').replace('June','junio').replace('July','julio').replace('August','agosto').replace('September','septiembre').replace('October','octubre').replace('November','noviembre').replace('December','diciembre')}", 11, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if bold and "PROJECT CHARTER" in text or text == "PROYECTO HAPPY JUMP":
            r.font.color.rgb = COLOR_PRIMARIO

    doc.add_page_break()

    # 1. Información general
    _titulo(doc, "1. INFORMACIÓN GENERAL DEL PROYECTO", 1)
    _tabla(doc, ["Campo", "Detalle"], [
        ["Nombre del proyecto", "Auditoría del Ciclo de Vida del Desarrollo de Software del Sistema Happy Jump"],
        ["Código del proyecto", "AUD-SDLC-HAPPY-JUMP-2026-001"],
        ["Versión del documento", "1.0"],
        ["Fecha de elaboración", fecha.strftime("%d de %B de %Y")],
        ["Fecha de inicio estimada", fecha.strftime("%d de %B de %Y")],
        ["Fecha de cierre estimada", "09 de septiembre de 2026"],
        ["Patrocinador", "Universidad Peruana Unión — Facultad de Ingeniería y Arquitectura"],
        ["Auditor líder", "[Nombre del auditor líder]"],
        ["Equipo auditor", "[Integrantes del equipo auditor]"],
        ["Clasificación", "Confidencial — Uso interno académico"],
        ["Sistema auditado", "Happy Jump"],
        ["Organización auditada", "Equipo desarrollador Happy Jump (GitHub: Randy236/apphappy)"],
    ])

    # 2. Antecedentes
    _titulo(doc, "2. ANTECEDENTES", 1)
    _titulo(doc, "2.1 Descripción del Sistema Happy Jump", 2)
    doc.add_paragraph(
        "Happy Jump es una aplicación móvil Android para la gestión de reservas de cancha "
        "deportiva (fútbol y vóley) y salones para eventos (cumpleaños y otros), desarrollada "
        "en el contexto del curso de Ingeniería de Software II de la Universidad Peruana Unión, "
        "bajo la supervisión del Ing. Rubén Roque Sucari."
    )
    doc.add_paragraph(
        "El sistema resuelve la necesidad de un centro de entretenimiento y eventos de registrar "
        "reservas por horarios, controlar pagos (completos o con adelanto), consultar disponibilidad "
        "mediante calendario, generar reportes de ingresos para el administrador y gestionar "
        "sesiones de personal con roles diferenciados (trabajador y administrador)."
    )

    _titulo(doc, "2.2 Arquitectura y Tecnología", 2)
    doc.add_paragraph(
        "Happy Jump adopta una arquitectura cliente-servidor de tres capas (no microservicios), "
        "documentada en el repositorio y en docs/DOCUMENTO_AVANCE_UNIDAD2.md:"
    )
    _bullets(doc, [
        "Aplicación móvil Android (Kotlin + Jetpack Compose) — módulos Cancha, Salones, Reportes y Perfil.",
        "API REST monolítica (Node.js 18+ + Express) — autenticación JWT, validaciones de dominio, Swagger/OpenAPI.",
        "Base de datos relacional MySQL 8 — esquema happy_jump, migraciones versionadas en server/migrations/.",
        "Integración continua: GitHub Actions (API Tests, SonarCloud, Snyk), Jenkins (Jenkinsfile), scripts PowerShell.",
        "Calidad y pruebas: JUnit/Allure (Android), node:test (API), Postman, k6, OWASP ZAP Desktop 2.17.",
        "Gestión ágil: Scrum + alineación CMMI Nivel 2 (docs/SCRUM_CMMI_HAPPY_JUMP.md, backlog Jira).",
    ])

    _titulo(doc, "2.3 Contexto de la Auditoría", 2)
    doc.add_paragraph(
        "La presente auditoría evalúa si el proceso de desarrollo de Happy Jump ha cumplido "
        "con los principios, metodologías y controles del ciclo de vida del software, conforme "
        "a ISO/IEC 12207, ISO/IEC 25010 y CMMI-DEV (Nivel 2 — Managed), considerando las "
        "evidencias en entregableunidad/, docs/ y el repositorio GitHub."
    )

    # 3. Justificación
    _titulo(doc, "3. JUSTIFICACIÓN", 1)
    _titulo(doc, "3.1 Necesidad de la Auditoría SDLC", 2)
    doc.add_paragraph(
        "El desarrollo de sistemas de información presenta altas tasas de fracaso cuando no "
        "se aplican controles formales en requisitos, diseño, pruebas y despliegue. Happy Jump "
        "maneja datos operativos sensibles (reservas, montos, PIN de usuarios) y requiere "
        "verificación de que las fases del SDLC fueron ejecutadas y documentadas."
    )
    _bullets(doc, [
        "El sistema opera con roles y sesión única; errores de seguridad impactan la operación diaria.",
        "El proyecto integra múltiples herramientas (Sonar, ZAP, Jenkins, Postman, k6) que deben trazarse al SDLC.",
        "El contexto académico exige evidencia de procesos formales (Scrum, CMMI, pruebas, documentación).",
        "Se requiere validar coherencia entre requisitos (REQ-XXX), casos de prueba (CP-XXX) e implementación.",
    ])

    _titulo(doc, "3.2 Beneficios de la Auditoría", 2)
    _bullets(doc, [
        "Identificación de brechas entre requisitos y entregables.",
        "Verificación del cumplimiento de estándares ISO/IEC 12207, ISO/IEC 25010 y CMMI-DEV.",
        "Reducción de riesgos antes de despliegue en producción real.",
        "Plan de acción correctiva basado en hallazgos objetivos.",
        "Línea base de calidad para futuras versiones del producto.",
    ])

    _titulo(doc, "3.3 Riesgos que se Pretenden Mitigar", 2)
    _bullets(doc, [
        "Incumplimiento de requisitos funcionales y no funcionales (REQ-001 a REQ-013).",
        "Vulnerabilidades en autenticación JWT, sesión única y control de roles.",
        "Inconsistencias entre arquitectura documentada e implementada.",
        "Ausencia o insuficiencia de evidencias de pruebas (unitarias, integración, E2E, seguridad, rendimiento).",
        "Falta de trazabilidad entre Jira, requisitos, casos CP-XXX y código en GitHub.",
    ])

    # 4. Problema u oportunidad
    _titulo(doc, "4. PROBLEMA U OPORTUNIDAD", 1)
    _titulo(doc, "4.1 Situación Identificada", 2)
    _bullets(doc, [
        "Verificar que la documentación de gestión (Scrum/Jira, cronograma, CMMI) esté completa y actualizada.",
        "Confirmar catálogo de requisitos (docs/REQUERIMIENTOS.csv) y casos de prueba (docs/CASOS_DE_PRUEBA.csv).",
        "Evaluar cobertura y calidad de pruebas en app/src/test/, server/test/, Postman, k6 y ZAP.",
        "Validar que el modelo SDLC (Scrum iterativo + entregables por unidad) esté explícito y aplicado.",
        "Revisar actas de cierre de sprints y entregables en entregableunidad/ (01 a 12).",
    ])
    _titulo(doc, "4.2 Oportunidad", 2)
    doc.add_paragraph(
        "La auditoría permite consolidar las buenas prácticas ya implementadas (CI/CD, SonarCloud ≥80%, "
        "informes Word automatizados, OpenAPI 100%) y corregir brechas antes de la evaluación final del curso."
    )

    # 5-6 Objetivos
    _titulo(doc, "5. OBJETIVO GENERAL", 1)
    doc.add_paragraph(
        "Realizar una auditoría formal del Ciclo de Vida del Desarrollo de Software del proyecto "
        "Happy Jump, evaluando procesos, controles, documentación, arquitectura y entregables, "
        "para determinar el cumplimiento de ISO/IEC 12207, ISO/IEC 25010, CMMI-DEV y buenas "
        "prácticas SDLC, identificar brechas y emitir recomendaciones de mejora."
    )

    _titulo(doc, "6. OBJETIVOS ESPECÍFICOS", 1)
    objetivos = [
        ("OE1", "Metodología y Gestión", "Evaluar planificación Scrum/Jira, sprints, Definition of Done y docs/SCRUM_CMMI_HAPPY_JUMP.md."),
        ("OE2", "Análisis de Requisitos", "Verificar docs/REQUERIMIENTOS.csv, trazabilidad con Jira y criterios de aceptación."),
        ("OE3", "Diseño del Sistema", "Evaluar coherencia Android + API + MySQL, OpenAPI (server/openapi.json), eliminado lógico y auditoría."),
        ("OE4", "Construcción y Codificación", "Revisar estándares Kotlin/JS, domain layer (server/src/domain/), GitHub, SonarCloud."),
        ("OE5", "Pruebas", "Verificar unitarias, integración Postman, E2E (CP-001–CP-022), seguridad ZAP, rendimiento k6."),
        ("OE6", "Implementación y Despliegue", "Evaluar Jenkinsfile, ci/, docker-compose, scripts de despliegue y manual de usuario."),
        ("OE7", "Mantenimiento y Soporte", "Verificar migraciones SQL, control de versiones Git, issues Jira y README."),
        ("OE8", "Documentación", "Evaluar manuales, guías en docs/, informes en entregableunidad/ y Swagger."),
        ("OE9", "Seguridad", "Verificar JWT, sesión única, pruebas Postman SEC y reporte OWASP ZAP DAST."),
        ("OE10", "Calidad del Producto", "Evaluar ISO/IEC 25010: funcionalidad, fiabilidad, usabilidad, eficiencia (k6), mantenibilidad (Sonar)."),
    ]
    for cod, nom, desc in objetivos:
        p = doc.add_paragraph()
        p.add_run(f"{cod} — {nom}: ").bold = True
        p.add_run(desc)

    # 7 Alcance
    _titulo(doc, "7. ALCANCE DE LA AUDITORÍA", 1)
    areas = [
        ("7.1 Gestión del Proyecto", "Scrum, Jira, sprints, retrospectivas, CMMI Nivel 2."),
        ("7.2 Requisitos", "REQ-001 a REQ-013, CASOS_DE_PRUEBA.csv CP-001 a CP-022."),
        ("7.3 Diseño y Arquitectura", "App Android, API Express, MySQL, OpenAPI, network security."),
        ("7.4 Construcción", "Repositorio GitHub, convenciones de commit, revisión de código."),
        ("7.5 Pruebas", "Unitarias JVM + Node, Postman, E2E manual, ZAP, k6, Allure, JaCoCo."),
        ("7.6 CI/CD y Calidad", "GitHub Actions, Jenkins, SonarCloud (cobertura ≥80%), Snyk."),
        ("7.7 Seguridad", "Autenticación, roles, OWASP ZAP, colección HappyJump-Seguridad."),
        ("7.8 Documentación y Entregables", "entregableunidad/entregable-01 … entregable-12, manual de usuario."),
    ]
    for tit, desc in areas:
        _titulo(doc, tit, 2)
        doc.add_paragraph(desc)

    doc.add_paragraph("Excluido del alcance: auditoría de infraestructura cloud en producción, pentesting externo profesional, certificación CMMI formal.")

    # 8 Metodología
    _titulo(doc, "8. METODOLOGÍA DE AUDITORÍA", 1)
    _tabla(doc, ["Fase", "Actividades", "Evidencias Happy Jump"], [
        ["Fase 1 — Planificación", "Criterios, plan, checklist SDLC", "Este Project Charter, Plan de Auditoría, Checklist"],
        ["Fase 2 — Revisión documental", "Docs, repo, entregables", "docs/, entregableunidad/, GitHub"],
        ["Fase 3 — Evaluación y reporte", "Matriz hallazgos, informe", "Matriz de Hallazgos, Informe Final"],
        ["Fase 4 — Seguimiento", "Plan correctivo, cierre", "Plan de Acción Correctiva, Acta de Cierre"],
    ])

    # 9 Entregables
    _titulo(doc, "9. ENTREGABLES DE LA AUDITORÍA", 1)
    _tabla(doc, ["N.º", "Entregable", "Descripción", "Fase"], [
        ["1", "Project Charter de Auditoría", "Inicio formal (este documento)", "Planificación"],
        ["2", "Plan de Auditoría", "Cronograma, recursos, procedimientos", "Fase 1"],
        ["3", "Checklist SDLC Adaptado", "Verificación por fase SDLC — Happy Jump", "Fase 1"],
        ["4", "Registro de Evidencias", "Compilación de evidencias del repo", "Fase 2"],
        ["5", "Papeles de Trabajo", "Procedimientos y resultados", "Fases 2 y 3"],
        ["6", "Matriz de Hallazgos", "Hallazgos, criterio, evidencia, impacto", "Fase 3"],
        ["7", "Matriz de Riesgos", "Riesgos derivados de hallazgos", "Fase 3"],
        ["8", "Informe Preliminar", "Borrador para revisión del auditado", "Fase 3"],
        ["9", "Informe Final de Auditoría", "Conclusiones y recomendaciones", "Fase 3"],
        ["10", "Plan de Acción Correctiva", "Acciones, responsables, plazos", "Fase 4"],
        ["11", "Acta de Cierre de Auditoría", "Cierre formal del proyecto de auditoría", "Fase 4"],
    ])

    # 10 Stakeholders auditoría
    _titulo(doc, "10. STAKEHOLDERS DE LA AUDITORÍA", 1)
    _tabla(doc, ["Stakeholder", "Rol", "Interés / Expectativa"], [
        ["Auditor líder", "Dirección", "Informe objetivo y fundamentado"],
        ["Equipo auditor", "Ejecución", "Evidencia suficiente y pertinente"],
        ["Patrocinador (UPU)", "Autoriza", "Calidad académica y profesional"],
        ["Ing. Rubén Roque Sucari", "Profesor / Supervisor", "Cumplimiento de estándares SDLC"],
        ["Equipo Happy Jump", "Auditado", "Retroalimentación y mejora del producto"],
    ])

    # 11 Stakeholders sistema
    _titulo(doc, "11. STAKEHOLDERS DEL SISTEMA HAPPY JUMP", 1)
    _tabla(doc, ["Stakeholder", "Descripción", "Interacción"], [
        ["Administrador", "Consulta reportes e ingresos", "App Android — pestaña Reportes"],
        ["Trabajador", "Registra reservas cancha y salones", "App Android — Cancha, Salones"],
        ["Administrador del sistema", "Mantiene API, BD y despliegue", "Servidor Node.js, MySQL, Jenkins"],
        ["Cliente final (indirecto)", "Persona que reserva", "Datos ingresados por trabajador"],
    ])

    # 12 Cronograma
    _titulo(doc, "12. CRONOGRAMA DE LA AUDITORÍA", 1)
    _tabla(doc, ["N.º", "Actividad / Hito", "Duración", "Responsable"], [
        ["1", "Inicio del proyecto de auditoría", "1 día", "Auditor líder"],
        ["2", "FASE 1 — Preparación y Planificación", "10 días háb.", "Equipo auditor"],
        ["2.1", "Criterios, objetivos y alcance", "2 días", "Auditor líder"],
        ["2.2", "Identificación SDLC Happy Jump (Scrum + CMMI)", "2 días", "Equipo auditor"],
        ["2.3", "Plan de Auditoría y Checklist SDLC", "5 días", "Equipo auditor"],
        ["3", "FASE 2 — Revisión documental", "15 días háb.", "Equipo auditor"],
        ["3.1", "Revisión docs/, entregableunidad/, GitHub", "5 días", "Equipo auditor"],
        ["3.2", "Entrevistas al equipo desarrollador", "4 días", "Equipo auditor"],
        ["3.3", "Revisión arquitectura, CI/CD, pruebas", "3 días", "Equipo auditor"],
        ["3.4", "Verificación evidencias Sonar, ZAP, k6, Postman", "3 días", "Equipo auditor"],
        ["4", "FASE 3 — Evaluación y reporte", "15 días háb.", "Equipo auditor"],
        ["4.1", "Matriz de Hallazgos y Riesgos", "7 días", "Equipo auditor"],
        ["4.2", "Informe Preliminar y Final", "8 días", "Auditor líder"],
        ["5", "FASE 4 — Seguimiento y cierre", "10 días háb.", "Equipo auditor"],
        ["6", "Cierre y archivo", "3 días", "Auditor líder"],
    ])

    # 13 Evidencias Happy Jump (específico del proyecto)
    _titulo(doc, "13. EVIDENCIAS DEL SISTEMA A AUDITAR (HAPPY JUMP)", 1)
    doc.add_paragraph("Referencia de evidencias ya producidas en el proyecto, sujetas a verificación:")
    _tabla(doc, ["Área SDLC", "Evidencia en el repositorio"], [
        ["Requisitos", "docs/REQUERIMIENTOS.csv, docs/CASOS_DE_PRUEBA.csv, Jira backlog"],
        ["Diseño / API", "server/openapi.json, Swagger UI /swagger-ui/"],
        ["Construcción", "app/ (Kotlin), server/ (Node.js), GitHub Randy236/apphappy"],
        ["Pruebas unitarias", "app/src/test/, server/test/, Allure, npm run test:ci"],
        ["Integración", "docs/postman/, scripts/postman_reports/"],
        ["E2E / Sistema", "entregable-06/07, CP-001–CP-022, manual capturas"],
        ["Seguridad", "entregable-08/, OWASP ZAP, Postman SEC-001–SEC-017"],
        ["Rendimiento", "entregable8.1/, k6/performance-api.js, INFORME_PRUEBAS_RENDIMIENTO.docx"],
        ["Calidad estática", "entregable-03/04/, SonarCloud ≥80%, Snyk, JaCoCo"],
        ["CI/CD", "Jenkinsfile, .github/workflows/, docs/GUIA_JENKINS_CI_CD.md"],
        ["Procesos", "docs/SCRUM_CMMI_HAPPY_JUMP.md, evidencias-scrum-cmmi/"],
        ["Usuario", "entregable-09/manualdeusuario.pdf, docs/informes-item9/"],
    ])

    # 14-20 secciones finales
    _titulo(doc, "14. CRITERIOS DE AUDITORÍA", 1)
    _bullets(doc, [
        "ISO/IEC 12207 — Procesos del ciclo de vida del software.",
        "ISO/IEC 25010 — Modelo de calidad del producto software.",
        "CMMI-DEV Nivel 2 — Áreas de proceso Managed (REQM, PP, PMC, SAM, MA, PPQA, CM).",
        "Definition of Done del proyecto (docs/SCRUM_CMMI_HAPPY_JUMP.md).",
        "Rúbrica y entregables del curso Ingeniería de Software II.",
    ])

    _titulo(doc, "15. CRITERIOS DE ACEPTACIÓN DEL PROYECTO DE AUDITORÍA", 1)
    _bullets(doc, [
        "Ejecución de todas las fases conforme al Plan de Auditoría.",
        "Evidencia suficiente en Matriz de Hallazgos.",
        "Informe Final elaborado, revisado y aprobado.",
        "Plan de Acción Correctiva presentado por el equipo auditado.",
        "Acta de Cierre firmada por Auditor líder y Patrocinador.",
    ])

    _titulo(doc, "16. SUPUESTOS", 1)
    _bullets(doc, [
        "El equipo Happy Jump colaborará entregando documentación y acceso al repositorio.",
        "La documentación en GitHub y entregableunidad/ es auténtica.",
        "El auditor tiene competencia en Android, Node.js y metodologías SDLC.",
        "La auditoría se realiza sobre el estado actual del sistema sin cambios sustanciales durante la ejecución.",
    ])

    _titulo(doc, "17. RESTRICCIONES", 1)
    _bullets(doc, [
        "Tiempo máximo: diez semanas (calendario académico).",
        "Recursos limitados del equipo auditor.",
        "Entorno local (Laragon/MySQL); puede no reflejar producción.",
        "Confidencialidad de credenciales demo (Admin/Rosisela PIN 1234).",
        "Independencia: el auditor no debe ser evaluador directo del mismo entregable simultáneamente.",
    ])

    _titulo(doc, "18. APROBACIONES", 1)
    doc.add_paragraph(
        "El presente Project Charter de Auditoría SDLC del Proyecto Happy Jump se somete "
        "a aprobación formal. La firma implica el inicio del proyecto de auditoría."
    )
    _tabla(doc, ["Rol", "Nombre", "Firma", "Fecha"], [
        ["Auditor líder", "", "", ""],
        ["Patrocinador (UPU)", "", "", ""],
        ["Representante equipo Happy Jump", "", "", ""],
        ["Ing. Rubén Roque Sucari (Supervisor)", "", "", ""],
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Documento de uso interno — Clasificación: Confidencial — Versión: 1.0 — "
        "Código: AUD-SDLC-HAPPY-JUMP-2026-001"
    )
    r.italic = True
    r.font.size = Pt(9)

    salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(salida)
    return salida
