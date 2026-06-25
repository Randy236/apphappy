"""
Motor educativo: ejecuta colecciones Postman v2.1 y genera informe Word empresarial.
Librerías: requests, python-docx (ver requirements.txt).
"""

from __future__ import annotations

import json
import re
import subprocess
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# Colores corporativos Happy Jump
COLOR_PRIMARIO = RGBColor(0x18, 0xB5, 0x8A)
COLOR_TEXTO = RGBColor(0x1E, 0x29, 0x3B)
COLOR_OK = RGBColor(0x16, 0xA3, 0x4A)
COLOR_FAIL = RGBColor(0xDC, 0x26, 0x26)
COLOR_HEADER_BG = "18B58A"
COLOR_HEADER_TEXT = "FFFFFF"


@dataclass
class CasoResultado:
    codigo: str
    nombre: str
    metodo: str
    endpoint: str
    restriccion: str
    esperado: str
    obtenido: int
    paso: bool
    detalle: str = ""
    es_setup: bool = False


@dataclass
class InformeEjecucion:
    titulo: str
    coleccion: str
    base_url: str
    fecha: datetime = field(default_factory=datetime.now)
    casos: list[CasoResultado] = field(default_factory=list)
    error_conexion: str | None = None

    @property
    def total(self) -> int:
        return len([c for c in self.casos if not c.es_setup or c.codigo.startswith(("VAL", "SEC"))])

    @property
    def pruebas(self) -> list[CasoResultado]:
        return [c for c in self.casos if c.codigo.startswith(("VAL", "SEC"))]

    @property
    def pasados(self) -> int:
        return sum(1 for c in self.pruebas if c.paso)

    @property
    def fallidos(self) -> int:
        return sum(1 for c in self.pruebas if not c.paso)


def _subst(text: str, variables: dict[str, str]) -> str:
    if not text:
        return text
    out = text
    for key, val in variables.items():
        out = out.replace(f"{{{{{key}}}}}", str(val))
    return out


def _url_from_request(req: dict, variables: dict[str, str]) -> str:
    raw = req.get("url")
    if isinstance(raw, str):
        return _subst(raw, variables)
    if isinstance(raw, dict):
        host = _subst(raw.get("raw") or "", variables)
        if host:
            return host
        protocol = raw.get("protocol", "http")
        host_parts = ".".join(
            p.get("value", "") for p in raw.get("host", []) if isinstance(p, dict)
        ) or "localhost"
        path = "/".join(
            p.get("value", "") for p in raw.get("path", []) if isinstance(p, dict)
        )
        return f"{protocol}://{host_parts}/{path}"
    return variables.get("baseUrl", "http://localhost:3000")


def _headers_from_request(req: dict, variables: dict[str, str]) -> dict[str, str]:
    headers: dict[str, str] = {}
    for h in req.get("header") or []:
        if h.get("disabled"):
            continue
        key = h.get("key", "")
        val = _subst(h.get("value", ""), variables)
        if key:
            headers[key] = val
    return headers


def _body_from_request(req: dict, variables: dict[str, str]) -> str | None:
    body = req.get("body") or {}
    if body.get("mode") == "raw":
        return _subst(body.get("raw", ""), variables)
    return None


def _parse_expected_statuses(exec_lines: list[str]) -> list[int] | None:
    text = "\n".join(exec_lines)
    statuses: list[int] = []
    for m in re.finditer(r"\.status\((\d+)\)", text):
        statuses.append(int(m.group(1)))
    m_one = re.search(r"oneOf\(\[([^\]]+)\]\)", text)
    if m_one:
        for part in m_one.group(1).split(","):
            part = part.strip()
            if part.isdigit():
                statuses.append(int(part))
    return statuses if statuses else None


def _codigo_caso(nombre: str) -> str:
    m = re.match(r"^(VAL|SEC)-\d+", nombre)
    if m:
        return m.group(0)
    return ""


def _es_setup(nombre: str) -> bool:
    n = nombre.lower()
    return "setup" in n or nombre.startswith("Login ") or nombre.startswith("Logout ")


def _restriccion_desde_nombre(codigo: str, nombre: str, tipo: str) -> str:
    mapa_val = {
        "VAL-001": "Nombre de usuario: solo letras, sin números",
        "VAL-002": "PIN: solo dígitos numéricos",
        "VAL-003": "PIN: máximo 5 dígitos",
        "VAL-004": "PIN: mínimo 4 dígitos",
        "VAL-005": "Reportes: periodo debe ser diario | semanal | mensual",
        "VAL-006": "Cancha: deporte Futbol o Voley únicamente",
        "VAL-007": "Cancha: montoTotal ≥ 0",
        "VAL-008": "Cancha: fecha formato YYYY-MM-DD",
        "VAL-009": "Cancha: adelanto ≤ montoTotal",
        "VAL-010": "Salones: nombre de salón de lista permitida",
        "VAL-011": "Cambio PIN: máximo 5 dígitos",
        "VAL-012": "Salón cumpleaños: nombreCumpleanero obligatorio",
    }
    mapa_sec = {
        "SEC-001": "Recurso protegido sin token → 401",
        "SEC-002": "Token JWT inválido → 401",
        "SEC-003": "Credenciales incorrectas → 401",
        "SEC-004": "Sesión única: segundo login → 409",
        "SEC-005": "Token invalidado tras logout → 401",
        "SEC-006": "Trabajador sin acceso a reportes → 403",
        "SEC-007": "Trabajador no cambia PIN ajeno → 403",
        "SEC-010": "Intento inyección SQL en login bloqueado",
        "SEC-016": "Health público accesible (riesgo aceptado dev)",
        "SEC-017": "PIN nuevo menor a 4 dígitos → 400",
    }
    if codigo in mapa_val:
        return mapa_val[codigo]
    if codigo in mapa_sec:
        return mapa_sec[codigo]
    if _es_setup(nombre):
        return "Preparación de sesión / token"
    return nombre


def _aplicar_variables_post_test(
    exec_lines: list[str], response: requests.Response, variables: dict[str, str]
) -> None:
    text = "\n".join(exec_lines)
    if response.status_code != 200:
        return
    try:
        data = response.json()
    except (json.JSONDecodeError, ValueError):
        return
    if "tokenAdmin" in text and "set('tokenAdmin'" in text.replace(" ", ""):
        if token := data.get("token"):
            variables["tokenAdmin"] = token
    if "tokenWorker" in text:
        if token := data.get("token"):
            variables["tokenWorker"] = token
    if "workerId" in text and "usuario.id" in text:
        usuario = data.get("usuario") or {}
        if uid := usuario.get("id"):
            variables["workerId"] = str(uid)


def limpiar_sesiones_api(server_root: Path) -> None:
    """Libera active_token en MySQL (igual que npm run k6:prepare)."""
    script = server_root / "scripts" / "clear-sessions-for-k6.mjs"
    if not script.is_file():
        return
    try:
        subprocess.run(
            ["node", str(script)],
            cwd=str(server_root),
            check=True,
            capture_output=True,
            text=True,
            timeout=15,
        )
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        pass


def ejecutar_coleccion(
    collection_path: Path, base_url: str, titulo: str, server_root: Path | None = None
) -> InformeEjecucion:
    with collection_path.open(encoding="utf-8") as f:
        coleccion = json.load(f)

    variables: dict[str, str] = {"baseUrl": base_url.rstrip("/")}
    for var in coleccion.get("variable") or []:
        if var.get("key"):
            variables[var["key"]] = str(var.get("value", ""))

    informe = InformeEjecucion(
        titulo=titulo,
        coleccion=collection_path.name,
        base_url=base_url,
    )

    # Verificar API
    try:
        requests.get(f"{base_url.rstrip('/')}/health", timeout=5)
    except requests.RequestException as exc:
        informe.error_conexion = str(exc)
        return informe

    if server_root:
        limpiar_sesiones_api(server_root)

    session = requests.Session()
    for item in coleccion.get("item") or []:
        nombre = item.get("name", "Sin nombre")
        req = item.get("request") or {}
        metodo = (req.get("method") or "GET").upper()
        url = _url_from_request(req, variables)
        headers = _headers_from_request(req, variables)
        body = _body_from_request(req, variables)

        exec_lines: list[str] = []
        for ev in item.get("event") or []:
            if ev.get("listen") == "test":
                exec_lines.extend(ev.get("script", {}).get("exec") or [])

        esperados = _parse_expected_statuses(exec_lines) or [200]
        codigo = _codigo_caso(nombre)
        es_setup = _es_setup(nombre) and not codigo

        try:
            kwargs: dict[str, Any] = {"headers": headers, "timeout": 15}
            if body is not None and metodo in ("POST", "PUT", "PATCH"):
                kwargs["data"] = body.encode("utf-8")
                if "Content-Type" not in headers:
                    kwargs["headers"] = {**headers, "Content-Type": "application/json"}
            resp = session.request(metodo, url, **kwargs)
        except requests.RequestException as exc:
            informe.casos.append(
                CasoResultado(
                    codigo=codigo or "—",
                    nombre=nombre,
                    metodo=metodo,
                    endpoint=url.replace(base_url, "") or url,
                    restriccion=_restriccion_desde_nombre(codigo, nombre, ""),
                    esperado=", ".join(str(s) for s in esperados),
                    obtenido=0,
                    paso=False,
                    detalle=f"Error de red: {exc}",
                    es_setup=es_setup,
                )
            )
            continue

        _aplicar_variables_post_test(exec_lines, resp, variables)
        paso = resp.status_code in esperados
        detalle = ""
        if not paso:
            try:
                detalle = resp.json().get("error", resp.text[:200])
            except (json.JSONDecodeError, ValueError):
                detalle = resp.text[:200]

        informe.casos.append(
            CasoResultado(
                codigo=codigo or ("SETUP" if es_setup else "—"),
                nombre=nombre,
                metodo=metodo,
                endpoint=url.replace(variables["baseUrl"], "") or "/",
                restriccion=_restriccion_desde_nombre(codigo, nombre, ""),
                esperado=", ".join(str(s) for s in esperados),
                obtenido=resp.status_code,
                paso=paso,
                detalle=str(detalle)[:120] if detalle else "",
                es_setup=es_setup,
            )
        )

    return informe


def _set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _estilo_celda_header(cell) -> None:
    _set_cell_shading(cell, COLOR_HEADER_BG)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)


def generar_word(informe: InformeEjecucion, salida: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # Portada
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("HAPPY JUMP")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_PRIMARIO

    subt = doc.add_paragraph()
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subt.add_run(informe.titulo)
    rs.bold = True
    rs.font.size = Pt(18)
    rs.font.color.rgb = COLOR_TEXTO

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Fecha de ejecución: {informe.fecha.strftime('%d/%m/%Y %H:%M')}\n").font.size = Pt(11)
    meta.add_run(f"API: {informe.base_url}\n").font.size = Pt(11)
    meta.add_run(f"Colección: {informe.coleccion}\n").font.size = Pt(11)
    meta.add_run("Proyecto: Sistema de reservas Happy Jump").font.size = Pt(11)

    doc.add_page_break()

    # Resumen ejecutivo
    h1 = doc.add_heading("1. Resumen ejecutivo", level=1)
    h1.runs[0].font.color.rgb = COLOR_PRIMARIO

    if informe.error_conexion:
        p = doc.add_paragraph()
        p.add_run("No se pudo conectar con la API. ").bold = True
        p.add_run(informe.error_conexion)
        doc.save(salida)
        return salida

    resumen = doc.add_table(rows=2, cols=4)
    resumen.alignment = WD_TABLE_ALIGNMENT.CENTER
    resumen.style = "Table Grid"
    headers = ["Total pruebas", "Aprobadas", "Fallidas", "% Éxito"]
    valores = [
        str(len(informe.pruebas)),
        str(informe.pasados),
        str(informe.fallidos),
        f"{(informe.pasados / len(informe.pruebas) * 100):.0f}%"
        if informe.pruebas
        else "—",
    ]
    for i, h in enumerate(headers):
        resumen.rows[0].cells[i].text = h
        _estilo_celda_header(resumen.rows[0].cells[i])
        resumen.rows[1].cells[i].text = valores[i]
        resumen.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(
        "Este informe documenta la ejecución automatizada de pruebas de integración "
        "equivalentes a la colección Postman, verificando que la API rechace entradas "
        "inválidas y cumpla las reglas de seguridad definidas."
    )

    # Tabla detalle
    doc.add_heading("2. Resultados por caso de prueba", level=1).runs[0].font.color.rgb = COLOR_PRIMARIO

    cols = [
        "Código",
        "Caso",
        "Método",
        "Endpoint",
        "Restricción / validación",
        "HTTP esperado",
        "HTTP obtenido",
        "Resultado",
    ]
    filas = [c for c in informe.casos if c.codigo.startswith(("VAL", "SEC"))]
    tabla = doc.add_table(rows=1 + len(filas), cols=len(cols))
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(cols):
        tabla.rows[0].cells[i].text = h
        _estilo_celda_header(tabla.rows[0].cells[i])

    for row_idx, caso in enumerate(filas, start=1):
        cells = tabla.rows[row_idx].cells
        datos = [
            caso.codigo,
            caso.nombre,
            caso.metodo,
            caso.endpoint[:40],
            caso.restriccion[:50],
            caso.esperado,
            str(caso.obtenido),
            "APROBADO" if caso.paso else "FALLIDO",
        ]
        for i, val in enumerate(datos):
            cells[i].text = val
            if i == 7:
                run = cells[i].paragraphs[0].runs[0] if cells[i].paragraphs[0].runs else cells[i].paragraphs[0].add_run(val)
                run.font.bold = True
                run.font.color.rgb = COLOR_OK if caso.paso else COLOR_FAIL

    # Restricciones del sistema
    doc.add_page_break()
    doc.add_heading("3. Reglas de validación en la API", level=1).runs[0].font.color.rgb = COLOR_PRIMARIO
    reglas = [
        ("Nombre usuario", "Solo letras, sin números, máx. 100 caracteres"),
        ("PIN", "Solo números, mín. 4 y máx. 5 dígitos"),
        ("nombreCliente", "Obligatorio, máx. 200 caracteres, caracteres permitidos"),
        ("Fecha", "Formato ISO YYYY-MM-DD"),
        ("Montos", "≥ 0, adelanto ≤ total"),
        ("Deporte", "Futbol o Voley"),
        ("Salón", "Lista cerrada de salones válidos"),
        ("Autenticación", "JWT obligatorio en rutas protegidas"),
    ]
    t2 = doc.add_table(rows=1 + len(reglas), cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Campo / aspecto"
    t2.rows[0].cells[1].text = "Regla"
    _estilo_celda_header(t2.rows[0].cells[0])
    _estilo_celda_header(t2.rows[0].cells[1])
    for i, (campo, regla) in enumerate(reglas, start=1):
        t2.rows[i].cells[0].text = campo
        t2.rows[i].cells[1].text = regla

    doc.add_paragraph()
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.add_run(
        f"Generado automáticamente — Happy Jump — {informe.fecha.strftime('%Y')}"
    ).font.size = Pt(9)
    pie.runs[0].font.italic = True

    salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(salida)
    return salida


def imprimir_resumen(informe: InformeEjecucion) -> int:
    print(f"\n{'=' * 60}")
    print(informe.titulo)
    print(f"API: {informe.base_url}")
    print(f"Colección: {informe.coleccion}")
    print(f"{'=' * 60}")
    if informe.error_conexion:
        print(f"ERROR: No hay conexión con la API — {informe.error_conexion}")
        print("Levanta MySQL y ejecuta: cd server && npm start")
        return 1
    for c in informe.casos:
        icono = "OK" if c.paso else "FAIL"
        print(f"  [{icono}] {c.nombre}: esperado {c.esperado}, obtenido {c.obtenido}")
    print(f"\nPruebas: {informe.pasados}/{len(informe.pruebas)} aprobadas")
    return 0 if informe.fallidos == 0 else 1
